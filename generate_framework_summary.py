from docx import Document
from pathlib import Path

summary = Document()
summary.add_heading('Pytest Playwright Framework Summary', level=1)
summary.add_paragraph(
    'This framework is built using pytest with Playwright for browser automation against OrangeHRM. '
    'It includes page objects, shared fixtures, reporting, and parallel execution support.'
)
summary.add_heading('Project Structure', level=2)
summary.add_paragraph('Root files:', style='List Bullet')
summary.add_paragraph('- conftest.py: shared fixtures and failure hooks', style='List Bullet')
summary.add_paragraph('- pytest.ini: pytest configuration including html reporting, retry, parallel execution, and Playwright video retention', style='List Bullet')
summary.add_paragraph('- requirements.txt: project dependencies and plugins', style='List Bullet')
summary.add_paragraph('- pages/: page object classes for application modules', style='List Bullet')
summary.add_paragraph('- tests/: test cases organized by feature', style='List Bullet')
summary.add_heading('Key Components', level=2)
summary.add_paragraph('Fixtures:', style='List Bullet')
summary.add_paragraph(
    'The launch_application fixture in conftest.py uses pytest-playwright fixtures playwright and browser_name '
    'to launch the browser and navigate to the application. It also manages browser context teardown after each test.',
    style='List Bullet'
)
summary.add_paragraph('Failure handling:', style='List Bullet')
summary.add_paragraph(
    'A pytest hook captures screenshots when tests fail and saves them to the screenshots/ directory.',
    style='List Bullet'
)
summary.add_paragraph('Reporting:', style='List Bullet')
summary.add_paragraph(
    'The framework generates an HTML report via pytest-html and is configured to retain videos on failure with --video=retain-on-failure.',
    style='List Bullet'
)
summary.add_heading('Page Object Model', level=2)
summary.add_paragraph(
    'The pages/ directory contains page objects that encapsulate UI actions and locators. Example classes:',
    style='List Bullet'
)
summary.add_paragraph('- BasePage: common wrappers for click, fill, text, visibility, and wait operations.', style='List Bullet')
summary.add_paragraph('- LoginPage: login actions and error handling.', style='List Bullet')
summary.add_paragraph('- DashboardPage: dashboard navigation and verification methods.', style='List Bullet')
summary.add_paragraph('- PimPage: employee workflow actions.', style='List Bullet')
summary.add_paragraph('- RecruitmentPage: recruitment search interactions.', style='List Bullet')
summary.add_heading('Test Design', level=2)
summary.add_paragraph(
    'Tests are located in the tests/ directory and use the launch_application fixture for browser setup. '
    'They are organized by feature and leverage pytest markers such as smoke and regression. Parametrized tests exist for data-driven login validation.'
)
summary.add_heading('Execution', level=2)
summary.add_paragraph('Recommended commands:')
summary.add_paragraph('pytest', style='List Number')
summary.add_paragraph('pytest --browsers=chromium,firefox,webkit', style='List Number')
summary.add_paragraph('pytest --collect-only -q', style='List Number')
summary.add_heading('Dependencies', level=2)
summary.add_paragraph(
    'Core dependencies include: pytest, pytest-playwright, pytest-html, pytest-xdist, pytest-rerunfailures, '
    'playwright, allure-pytest.'
)
summary.add_heading('Improvements', level=2)
summary.add_paragraph('Potential next improvements:', style='List Bullet')
summary.add_paragraph('- Centralize artifacts in a single artifacts/ folder.', style='List Bullet')
summary.add_paragraph('- Add a CI workflow for automated test runs.', style='List Bullet')
summary.add_paragraph('- Add video and screenshot embedding into the HTML report.', style='List Bullet')
summary.add_paragraph('- Introduce linting and formatting with black, flake8, and pre-commit.', style='List Bullet')

path = Path('framework_summary.docx')
summary.save(path)
print(path.resolve())
